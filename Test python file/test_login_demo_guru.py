from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
import time
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

# Headless mode тохируулах
options = Options()
options.add_argument("--headless")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--disable-gpu")

# 1. Он сартай, давхцахгүй фолдер үүсгэх
now = datetime.now().strftime("%Y%m%d_%H%M%S")
base_folder = f"test_result_{now}"
os.makedirs(base_folder, exist_ok=True)
img_folder = os.path.join(base_folder, "screenshots")
os.makedirs(img_folder, exist_ok=True)

# 2. Word файл үүсгэх
doc = Document()
doc.add_heading('Guru99 Bank Login Test Report', 0)

print("Browser нээгдэж байна...")
driver = webdriver.Chrome(options=options)
driver.get("https://demo.guru99.com/V4/")
time.sleep(1)
img1 = os.path.join(img_folder, "01_open_browser.png")
driver.save_screenshot(img1)
doc.add_paragraph("1. Guru99 Bank login хуудас ачааллаа.")
doc.add_picture(img1, width=Inches(4))

print("Нэвтрэх мэдээлэл оруулж байна...")
driver.find_element(By.NAME, "uid").send_keys("mngr625959")
driver.find_element(By.NAME, "password").send_keys("AdUryqy")
img2 = os.path.join(img_folder, "02_enter_credentials.png")
driver.save_screenshot(img2)
doc.add_paragraph("2. Username болон Password орууллаа.")
doc.add_picture(img2, width=Inches(4))

driver.find_element(By.NAME, "btnLogin").click()
time.sleep(2)
img3 = os.path.join(img_folder, "03_after_login.png")
driver.save_screenshot(img3)
doc.add_paragraph("3. Login дарсны дараах хуудасны screenshot.")
doc.add_picture(img3, width=Inches(4))

print("Dashboard шалгаж байна...")
# Амжилттай login бол "Manger Id : mngr34926" гэсэн текст гарна
assert "Manger Id : mngr625959" in driver.page_source
doc.add_paragraph("4. 'Manger Id : mngr34926' гарчиг амжилттай илэрлээ.")

print("Тест амжилттай дууслаа!")
driver.quit()
doc.add_paragraph("5. Тест амжилттай дууслаа.")

# 3. Word файлыг хадгалах
word_path = os.path.join(base_folder, "test_report.docx")
doc.save(word_path)

print(f"Бүх screenshot болон тайлан: {base_folder} хавтсанд хадгалагдлаа.")
print(f"Word файл: {word_path} хадгалагдлаа.")
print("Тест дууслаа!")