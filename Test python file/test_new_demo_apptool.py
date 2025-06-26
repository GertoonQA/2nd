from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import NoSuchElementException, TimeoutException
from docx import Document
from docx.shared import Inches
from datetime import datetime
import time
import os


def setup_driver(headless=False):
    options = Options()
    if headless:
        options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    driver = webdriver.Chrome(options=options)
    driver.set_window_size(1280, 1024)  # Илүү тогтвортой харуулна
    return driver

def create_result_folders():
    base_dir = r"D:\\2nd project\\results"
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_folder = os.path.join(base_dir, f"applitools_test_{now}")
    img_folder = os.path.join(base_folder, "screenshots")
    os.makedirs(img_folder, exist_ok=True)
    return base_folder, img_folder

def add_screenshot(doc, driver, img_folder, filename, description):
    try:
        filepath = os.path.join(img_folder, filename)
        driver.save_screenshot(filepath)
        doc.add_paragraph(description)
        doc.add_picture(filepath, width=Inches(4))
    except Exception as e:
        doc.add_paragraph(f"⚠️ Screenshot авахад алдаа гарлаа: {e}")

def wait_for_element(driver, by, locator, timeout=10):
    """ Элемент гарах хүртэл хүлээх функц """
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC

    try:
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located((by, locator)))
        return element
    except TimeoutException:
        return None

def main():
    base_folder, img_folder = create_result_folders()
    doc = Document()
    doc.add_heading('Applitools Demo Site Test Report', 0)

    driver = setup_driver(headless=False)
    driver.get("https://demo.applitools.com/")
    time.sleep(1)  # Хуудас бүрэн ачаалагдахад бага зэрэг хүлээнэ

    add_screenshot(doc, driver, img_folder, "01_open_browser.png", "1. Applitools demo сайт нээгдлээ.")

    # --- Login Test ---
    print("Login хийгдэж байна...")
    try:
        username = wait_for_element(driver, By.ID, "username")
        password = wait_for_element(driver, By.ID, "password")
        login_btn = wait_for_element(driver, By.ID, "log-in")

        if not all([username, password, login_btn]):
            raise Exception("Login элемент олдсонгүй")

        username.clear()
        username.send_keys("demo_user")
        password.clear()
        password.send_keys("demo_pass")

        add_screenshot(doc, driver, img_folder, "02_before_login.png", "2. Login мэдээлэл оруулсан байдал.")

        login_btn.click()
        time.sleep(3)  # Нэвтрэлтийн дараах хуудсыг ачаалахад хүлээнэ

        add_screenshot(doc, driver, img_folder, "03_after_login.png", "3. Login хийсний дараах хуудас.")

        # Амжилттай нэвтэрсэн эсэхийг шалгах (dashboard дээрх ямар нэг текстээр)
        assert "Your nearest branch closes in:" in driver.page_source
        doc.add_paragraph("Login амжилттай боллоо.")

    except Exception as e:
        print("⚠️ Login туршилт амжилтгүй:", e)
        doc.add_paragraph(f"⚠️ Login туршилт амжилтгүй: {e}")

    # --- Transfer Funds Test ---
    print("Transfer Funds туршиж байна...")
    try:
        transfer_tab = wait_for_element(driver, By.ID, "make-transaction")
        if not transfer_tab:
            raise Exception("Transfer Funds таб олдсонгүй")

        transfer_tab.click()
        time.sleep(2)

        add_screenshot(doc, driver, img_folder, "04_transfer_funds_tab.png", "4. Transfer Funds таб нээгдлээ.")

        amount_field = wait_for_element(driver, By.ID, "amount")
        recipient_field = wait_for_element(driver, By.ID, "recipient")
        transfer_btn = wait_for_element(driver, By.ID, "transfer-btn")

        if not all([amount_field, recipient_field, transfer_btn]):
            raise Exception("Transfer Funds формаас элемент олдсонгүй")

        amount_field.clear()
        amount_field.send_keys("100")
        recipient_field.clear()
        recipient_field.send_keys("Test Recipient")

        add_screenshot(doc, driver, img_folder, "05_filled_transfer_form.png", "5. Transfer Funds форма бөглөв.")

        transfer_btn.click()
        time.sleep(3)

        add_screenshot(doc, driver, img_folder, "06_transfer_result.png", "6. Transfer Funds үр дүнгийн хуудас.")

        # Үр дүнгийн текстээр шалгах
        assert "Transfer Complete!" in driver.page_source
        doc.add_paragraph("Transfer Funds амжилттай боллоо.")

    except Exception as e:
        print("⚠️ Transfer Funds туршилт амжилтгүй:", e)
        doc.add_paragraph(f"⚠️ Transfer Funds туршилт амжилтгүй: {e}")

    # --- Logout Test ---
    print("Logout хийж байна...")
    try:
        logout_btn = wait_for_element(driver, By.ID, "logout")
        if logout_btn:
            logout_btn.click()
            time.sleep(2)
            add_screenshot(doc, driver, img_folder, "07_logout.png", "7. Logout хийсэн байдал.")
            doc.add_paragraph("Logout амжилттай боллоо.")
        else:
            raise Exception("Logout товч олдсонгүй")
    except Exception as e:
        print("⚠️ Logout туршилт амжилтгүй:", e)
        doc.add_paragraph(f"⚠️ Logout туршилт амжилтгүй: {e}")

    driver.quit()
    doc.add_paragraph("Тест амжилттай дууслаа.")

    report_path = os.path.join(base_folder, "applitools_test_report.docx")
    doc.save(report_path)

    print(f"✅ Тестийн тайлан амжилттай хадгалагдлаа: {report_path}")

if __name__ == "__main__":
    main()
