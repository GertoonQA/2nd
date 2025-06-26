from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import ElementClickInterceptedException, NoAlertPresentException, UnexpectedAlertPresentException
from docx import Document
from docx.shared import Inches
from datetime import datetime
import time
import os

def setup_driver():
    options = Options()
    options.add_argument("--headless")  # Хэрвээ хүсвэл headless-г унтрааж болно
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    return webdriver.Chrome(options=options)

def create_result_folders():
    results_root = r"D:\\2nd project\\results"
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_folder = os.path.join(results_root, f"test_result_{now}")
    img_folder = os.path.join(base_folder, "screenshots")
    os.makedirs(img_folder, exist_ok=True)
    return base_folder, img_folder

def add_screenshot(doc, driver, img_folder, filename, description):
    try:
        path = os.path.join(img_folder, filename)
        driver.save_screenshot(path)
        doc.add_paragraph(description)
        doc.add_picture(path, width=Inches(4))
    except Exception as e:
        doc.add_paragraph(f"⚠️ Screenshot авахад алдаа гарлаа: {str(e)}")

def remove_obstructive_elements(driver):
    # Заримдаа click хийхэд саад болдог зургуудыг устгана
    for _ in range(10):
        removed = driver.execute_script("""
            var imgs = document.querySelectorAll('img');
            var count = imgs.length;
            imgs.forEach(function(img){ img.remove(); });
            return count;
        """)
        if removed == 0:
            break
        time.sleep(0.2)

def safe_click_by_link_text(driver, text):
    # Link text-ээр click хийх үед саадтай бол устгаж дахин оролдоно
    for _ in range(5):
        try:
            elem = driver.find_element(By.LINK_TEXT, text)
            if elem.is_displayed() and elem.is_enabled():
                elem.click()
                return True
        except ElementClickInterceptedException:
            print(f"Click хийхэд саад боллоо, дахин оролдож байна: {text}")
            remove_obstructive_elements(driver)
            time.sleep(0.2)
        except Exception as e:
            print(f"Click хийхэд алдаа: {e}")
            time.sleep(0.2)
    return False

def handle_unexpected_alert(driver, doc=None):
    try:
        alert = driver.switch_to.alert
        alert_text = alert.text
        print("⚠️ Alert илэрлээ:", alert_text)
        alert.accept()
        if doc:
            doc.add_paragraph(f"⚠️ Alert илэрсэн: {alert_text}")
        time.sleep(1)
        return True
    except NoAlertPresentException:
        return False

# === Main Test ===
if __name__ == "__main__":
    base_folder, img_folder = create_result_folders()
    doc = Document()
    doc.add_heading('Guru99 Bank Full Test Report', 0)

    driver = setup_driver()
    driver.get("https://demo.guru99.com/V4/")
    time.sleep(1)
    add_screenshot(doc, driver, img_folder, "01_open_browser.png", "1. Login хуудас нээгдлээ.")

    print("Нэвтрэх...")
    driver.find_element(By.NAME, "uid").send_keys("mngr625959")
    driver.find_element(By.NAME, "password").send_keys("AdUryqy")
    add_screenshot(doc, driver, img_folder, "02_credentials.png", "2. Username/Password оруулав.")

    driver.find_element(By.NAME, "btnLogin").click()
    time.sleep(2)
    add_screenshot(doc, driver, img_folder, "03_dashboard.png", "3. Login хийсний дараах хуудас.")

    assert "Manger Id : mngr625959" in driver.page_source
    doc.add_paragraph("4. Manger Id амжилттай шалгалаа.")

    # --- Balance Enquiry ---
    print("Balance Enquiry цэс рүү орж байна...")
    remove_obstructive_elements(driver)
    if safe_click_by_link_text(driver, "Balance Enquiry"):
        time.sleep(1)
        add_screenshot(doc, driver, img_folder, "04_balance_enquiry.png", "5. Balance Enquiry цэс рүү орсон байдал.")
    else:
        print("Balance Enquiry рүү орох боломжгүй байна.")

    # --- Fund Transfer ---
    print("Fund Transfer цэс рүү орж байна...")
    try:
        if safe_click_by_link_text(driver, "Fund Transfer"):
            time.sleep(1)

            driver.find_element(By.NAME, "payersaccount").send_keys("3308")
            driver.find_element(By.NAME, "payeeaccount").send_keys("3309")
            driver.find_element(By.NAME, "ammount").send_keys("100")
            driver.find_element(By.NAME, "desc").send_keys("Test transfer")

            img5 = os.path.join(img_folder, "05_fund_transfer_form.png")
            driver.save_screenshot(img5)
            doc.add_paragraph("6. Fund Transfer form бөглөв.")
            doc.add_picture(img5, width=Inches(4))

            driver.find_element(By.NAME, "AccSubmit").click()
            time.sleep(2)

            if handle_unexpected_alert(driver, doc):
                print("Alert гараад хаагдлаа, Fund Transfer үр дүнгийн зураг авахгүй.")
            else:
                img6 = os.path.join(img_folder, "06_fund_transfer_result.png")
                driver.save_screenshot(img6)
                doc.add_paragraph("7. Fund Transfer үр дүнгийн screenshot.")
                doc.add_picture(img6, width=Inches(4))
        else:
            print("Fund Transfer цэс рүү орох боломжгүй байна.")
    except Exception as e:
        print("⚠️ Fund Transfer хэсэгт алдаа:", e)
        doc.add_paragraph(f"⚠️ Fund Transfer хэсэгт алдаа: {str(e)}")

    # --- Mini Statement ---
    print("Mini Statement рүү орж байна...")
    if safe_click_by_link_text(driver, "Mini Statement"):
        time.sleep(1)
        driver.find_element(By.NAME, "accountno").send_keys("3308")
        driver.find_element(By.NAME, "AccSubmit").click()
        time.sleep(1)
        add_screenshot(doc, driver, img_folder, "07_mini_statement.png", "8. Mini Statement харагдсан байдал.")
    else:
        print("Mini Statement цэс рүү орох боломжгүй байна.")

    # --- Logout ---
    print("Log out хийж байна...")
    if safe_click_by_link_text(driver, "Log out"):
        time.sleep(1)
        add_screenshot(doc, driver, img_folder, "08_logout.png", "9. Амжилттай Log out хийв.")
    else:
        print("Log out хийх боломжгүй байна.")

    driver.quit()
    doc.add_paragraph("10. Тест амжилттай дууслаа.")

    word_path = os.path.join(base_folder, "test_report.docx")
    doc.save(word_path)

    print(f"✅ Бүх screenshot болон тайлан: {base_folder}")
    print(f"📄 Word файл: {word_path}")
    print("Тест амжилттай гүйцэтгэлээ.")
    print("Баярлалаа!")
    print("Сайн сайхан өдөр!")